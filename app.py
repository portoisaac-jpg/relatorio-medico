from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
from datetime import datetime

app = Flask(__name__)

# ── Mapeamentos ────────────────────────────────────────────────────────────────

DIAGNOSTICOS_LABELS = {
    "tea_1": "Transtorno do Espectro Autista (TEA) — nível 1 (leve)",
    "tea_2": "Transtorno do Espectro Autista (TEA) — nível 2 (moderado)",
    "tea_3": "Transtorno do Espectro Autista (TEA) — nível 3 (severo)",
    "tea_linguagem": "TEA com comprometimento da linguagem",
    "tdah": "Transtorno do Déficit de Atenção com Hiperatividade (TDAH)",
    "integracao_sens": "Transtorno de Integração Sensorial",
    "epilepsia": "Epilepsia",
    "epilepsia_refr": "Epilepsia refratária / de difícil controle",
    "atraso_neuro": "Atraso de neurodesenvolvimento",
    "atraso_fala": "Atraso de fala",
    "def_intelect": "Deficiência intelectual",
    "disturbio_sono": "Transtorno do sono",
    "ansiedade": "Transtorno de ansiedade",
    "fibromialgia": "Fibromialgia",
}

SINTOMAS_LABELS = {
    "choro_intenso": "Choro intenso",
    "nao_fixacao_olhar": "Não fixação do olhar",
    "ponta_pes": "Andar na ponta dos pés",
    "seletividade_alim": "Seletividade alimentar",
    "nao_aceita_comando": "Não aceita comando",
    "sem_nocao_perigo": "Não tem noção do perigo",
    "agressividade": "Agressividade",
    "atraso_fala_sint": "Atraso da fala",
    "atraso_cognitivo": "Atraso de desenvolvimento cognitivo",
    "hiperatividade": "Hiperatividade",
    "auto_mutilacao": "Auto-mutilação",
    "impulsividade": "Impulsividade",
    "dif_socializacao": "Dificuldade de socialização com outras crianças",
    "disturbio_sono_sint": "Distúrbios do sono",
    "estereotipias": "Estereotipias",
    "nao_verbal": "Paciente não verbal",
    "dif_concentracao": "Dificuldade de concentração",
    "dif_aprendizado": "Dificuldade de aprendizado escolar",
    "baixo_lim_frust": "Baixo limiar à frustração",
    "desatencao": "Desatenção",
    "procrastinacao": "Tendência à procrastinação",
}

FIBRO_SINTOMAS_LABELS = {
    "dor_difusa": "Dor muscular difusa generalizada",
    "fadiga_cronica": "Fadiga crônica",
    "sono_nao_reparador": "Sono não reparador",
    "fibro_fog": "Névoa mental (fibro fog) — dificuldade de memória e concentração",
    "cefaleia": "Cefaleia recorrente",
    "parestesia": "Parestesias (formigamentos)",
    "sii": "Síndrome do intestino irritável",
    "hipersensibilidade": "Hipersensibilidade a estímulos (luz, som, toque)",
    "ansiedade_fibro": "Ansiedade associada",
    "depressao_fibro": "Depressão associada",
    "rigidez_matinal": "Rigidez matinal",
}

FIBRO_CRITERIOS_LABELS = {
    "acr_1990": "ACR 1990 (pontos dolorosos)",
    "acr_2010": "ACR 2010/2011 (WPI + SS)",
    "acr_2016": "ACR 2016 (revisado)",
}

FIBRO_INTENSIDADE_LABELS = {
    "leve": "Leve (EVA 1–3)",
    "moderada": "Moderada (EVA 4–6)",
    "intensa": "Intensa (EVA 7–10)",
}

MEDS = [
    ("lamotrigina",   "Lamotrigina"),
    ("neozine",       "Neozine"),
    ("risperidona",   "Risperidona"),
    ("levetiracetam", "Levetiracetam"),
    ("neuliptil",     "Neuliptil"),
    ("depakene",      "Depakene (Valproato)"),
    ("frisium",       "Frisium (Clobazam)"),
    ("aristab",       "Aristab (Aripiprazol)"),
    ("ritalina",      "Ritalina (Metilfenidato)"),
    ("venvanse",      "Venvanse (Lisdexanfetamina)"),
    ("fluoxetina",    "Fluoxetina"),
    ("sertralina",    "Sertralina"),
    ("clonazepam",    "Clonazepam"),
    ("haloperidol",   "Haloperidol"),
    ("topiramato",    "Topiramato"),
    ("carbamazepina", "Carbamazepina"),
    ("fenobarbital",  "Fenobarbital"),
]

MED_RESPOSTA_LABELS = {
    "sem_efeito":    "sem efeito terapêutico",
    "efeito_parcial":"efeito parcial",
    "intoxicacao":   "causou intoxicação",
    "em_uso":        "em uso atualmente",
    "suspenso":      "suspenso por efeito adverso",
}

TERAPIAS_LABELS = {
    "terapia_ocupacional": "Terapia Ocupacional",
    "fonoaudiologia":      "Fonoaudiologia",
    "fisioterapia":        "Fisioterapia",
    "psicologia":          "Psicologia",
    "musicoterapia":       "Musicoterapia",
    "psicomotricidade":    "Psicomotricidade",
    "nutricao":            "Nutrição",
    "equoterapia":         "Equoterapia",
    "aba":                 "Análise do Comportamento Aplicada (ABA)",
    "hidroterapia":        "Hidroterapia",
}

JUSTIFICATIVAS_TEXTOS = {
    "falha_terapeutica": (
        "Diante da falha terapêutica na remissão dos sintomas do paciente com a terapia clássica, "
        "foi discutido com os responsáveis a possibilidade da modulação do sistema endocanabinóide "
        "a partir de produtos da Cannabis sativa. A descoberta deste sistema e seus receptores trouxe "
        "luz ao tratamento de diversas doenças crônicas com a possibilidade de menos efeitos adversos "
        "e melhora da saúde global de pacientes e cuidadores."
    ),
    "cfm_2113": (
        "Segundo a Resolução 2.113/14 do Conselho Federal de Medicina, protocolo de utilização do CBD "
        "para epilepsia refratária: o CBD deverá ser utilizado em adição às medicações que o paciente "
        "vinha utilizando anteriormente. O tratamento com CBD pode começar com 2,5mg/kg/dia por via oral, "
        "divididas em duas doses diárias, podendo ser aumentada quinzenalmente até a dose máxima de "
        "25mg/kg/dia, a fim de determinar a dose ideal com garantia de segurança e tolerabilidade."
    ),
    "cbd_thc": (
        "No TEA e no TDAH, o tratamento combinado de CBD e THC vem se mostrando bastante promissor devido "
        "à sua capacidade polifarmacológica, na qual sintomas como ansiedade, comportamentos estereotipados, "
        "insônia e epilepsias são tratados com um único produto. Além disso, é relatada melhora na aceitação "
        "de alimentos e apetite. A Cannabis conta ainda com mais de uma centena de canabinoides ativos "
        "(CBN, CBG, THCV etc.) e outros grupos de moléculas com efeitos farmacológicos relevantes, como "
        "flavonoides e terpenos, que se combinam garantindo um amplo efeito terapêutico."
    ),
    "full_spectrum": (
        "Foi escolhido o uso de óleo de Cannabis Full Spectrum por se tratar de produto integral, que garante "
        "além do CBD outros canabinoides (como o THC, com papel importante em sintomas motores e elétricos "
        "como as epilepsias), terpenos e flavonoides, com ações farmacológicas relevantes para o tratamento "
        "combinado. A literatura aponta que o uso integral geralmente apresenta ação farmacológica efetiva "
        "com doses menores, ação em sintomas secundários como apetite e ansiedade, além de menores efeitos "
        "adversos durante o uso."
    ),
    "neuroplasticidade": (
        "Estudos indicam que a neuroplasticidade é maior nos primeiros anos de vida, o que implica em uma "
        "resposta terapêutica mais efetiva. Portanto, o tratamento deve ser iniciado imediatamente e mantido "
        "por tempo indeterminado, sem interrupções."
    ),
    "bjhr_2023": (
        "Em junho de 2023 a Brazilian Journal of Health Review realizou revisão integrativa em que se concluiu "
        "que o uso do canabidiol no TEA possui relação com melhora na qualidade de vida dos pacientes — "
        "comportamento, hiperatividade, estereotipias, distúrbios do sono, comorbidades e convulsões — além "
        "de reduzir ansiedade, agressividade, inquietação e agitação."
    ),
    "epilepsia_dificil": (
        "A epilepsia de difícil controle é caracterizada por falha de controle de crises após tentativa de "
        "mais de dois anticonvulsivantes de primeira linha apropriados à idade e ao tipo de crise. O uso "
        "adequado de medicamentos auxiliares ajuda a controlar crises e minimizar fatores de risco e "
        "complicações, incluindo o status epilepticus, condição especialmente perigosa pelo risco de danos "
        "cerebrais permanentes."
    ),
    "epilepsia_combinacao": (
        "Nos casos de pacientes portadores de crises epilépticas e convulsivas, a formulação do tratamento "
        "à base de canabidiol deverá, em regra, ser associada à composição de outros canabinoides, a fim de "
        "que seja garantido maior potencial efetivo, jamais possível com o uso de fármacos isolados."
    ),
    "tea_sem_tratamento_especifico": (
        "O Transtorno do Espectro Autista (TEA) é uma comorbidade para a qual não há tratamento específico. "
        "Trata-se de transtorno neuropsiquiátrico desenvolvido na infância, manejado por equipe multiprofissional, "
        "com diversas terapêuticas adjuvantes — entre as quais o canabidiol, com objetivo de alívio de sinais "
        "específicos."
    ),
}

URGENCIAS_TEXTOS = {
    "urgencia_tea_tdah": (
        "Por se tratar de paciente dentro do TEA e/ou TDAH, há urgência em iniciar o tratamento com canabidiol, "
        "que se mostra eficaz na redução dos sintomas dessas patologias. Somado a isso, terapias multidisciplinares "
        "com psicólogo, fonoaudiólogo e terapeuta ocupacional são abordagem complementar essencial para o "
        "desenvolvimento de competências e habilidades para a vida adulta."
    ),
    "risco_status_epilept": (
        "Há risco aumentado de status epilepticus, condição especialmente perigosa pela possibilidade de causar "
        "danos cerebrais permanentes. O início imediato do tratamento é fundamental para controle adequado das crises."
    ),
    "risco_regressao": (
        "A não continuação do tratamento em curso poderá acarretar regressão dos avanços já alcançados pelo "
        "paciente, comprometendo seu desenvolvimento neurológico e funcional."
    ),
    "uso_continuo": (
        "O uso do óleo de Cannabis deve ser contínuo. Pode haver necessidade de aumento da concentração do "
        "medicamento com o passar do tempo; entretanto, a substituição por outros óleos de canabidiol sem "
        "avaliação médica prévia não é segura."
    ),
    "perda_aprendizado": (
        "A persistência das crises e dos sintomas vem causando prejuízos significativos no aprendizado e na "
        "socialização do paciente, sendo urgente o início do tratamento para evitar maior comprometimento."
    ),
}

PRODUTOS_LABELS = {
    "terramed_3000": "TERRAMED Oil Full Spectrum 3000mg (Fabricante TERRAMED/EUA)",
    "terramed_1500": "TERRAMED Oil Full Spectrum 1500mg (Fabricante TERRAMED/EUA)",
    "terramed_6000": "TERRAMED Oil Full Spectrum 6000mg (Fabricante TERRAMED/EUA)",
    "hempmeds_3000": "HempMeds CBD Real Scientific Hemp Oil 3000mg",
    "entourage_2000": "Entourage Full Spectrum 2000mg",
    "abrace_3000":   "Abrace Esperança Full Spectrum 3000mg (Brasil)",
    "prati_200":     "Prati-Donaduzzi Canabidiol 200mg/ml (Brasil)",
}

ACOMP_LABELS = {
    "equipe_multi":  "Equipe multidisciplinar especializada",
    "neurologista":  "Neurologista",
    "psiquiatra":    "Psiquiatra",
    "pediatra":      "Pediatra",
    "medico_assist": "Médico assistente prescritor",
}

PERIOD_LABELS = {
    "mensal":               "mensal",
    "bimestral":            "bimestral",
    "trimestral":           "trimestral",
    "semestral":            "semestral",
    "conforme_necessidade": "conforme necessidade",
}

# ── Helpers docx ──────────────────────────────────────────────────────────────

def set_font(run, name="Arial", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=13, bold=True, color=(31, 73, 125))
    else:
        set_font(run, size=11, bold=True, color=(54, 95, 145))
    return p

def add_body(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=10.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        rb = p.add_run(bold_prefix + " ")
        set_font(rb, size=10.5, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── Geração do documento ──────────────────────────────────────────────────────

def gerar_docx(f, logo_bytes=None):
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Logo ──
    if logo_bytes:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        img_stream = io.BytesIO(logo_bytes)
        run.add_picture(img_stream, width=Inches(2))
        doc.add_paragraph()

    # ── Título ──
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = titulo.add_run("RELATÓRIO MÉDICO")
    set_font(rt, size=16, bold=True, color=(31, 73, 125))

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = subtitulo.add_run("Prescrição de Canabidiol — Cannabis sativa")
    set_font(rs, size=11, color=(89, 89, 89))
    doc.add_paragraph()

    # ── Dados do médico ──
    medico      = f.get("medico", "")
    crm         = f.get("crm", "")
    especialidade = f.get("especialidade", "")
    med_line = f"Dr(a). {medico}  |  CRM {crm}"
    if especialidade:
        med_line += f"  |  {especialidade}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = p.add_run(med_line)
    set_font(rm, size=10, color=(89, 89, 89))

    add_divider(doc)
    doc.add_paragraph()

    # ════════ 1. IDENTIFICAÇÃO ════════
    add_heading(doc, "1. Identificação do Paciente")

    nome     = f.get("paciente_nome", "")
    idade    = f.get("paciente_idade", "")
    sexo_v   = f.get("paciente_sexo", "")
    sexo     = "Masculino" if sexo_v == "masculino" else ("Feminino" if sexo_v == "feminino" else "")
    rg       = f.get("paciente_rg", "")
    cpf      = f.get("paciente_cpf", "")
    endereco = f.get("paciente_endereco", "")
    filiacao = f.get("paciente_filiacao", "")

    if nome:      add_bullet(doc, nome, "Nome:")
    if idade:     add_bullet(doc, idade, "Idade:")
    if sexo:      add_bullet(doc, sexo, "Sexo:")
    if rg:        add_bullet(doc, rg, "RG:")
    if cpf:       add_bullet(doc, cpf, "CPF:")
    if endereco:  add_bullet(doc, endereco, "Endereço:")
    if filiacao:  add_bullet(doc, filiacao, "Filiação:")

    # ════════ 2. HISTÓRICO CLÍNICO ════════
    add_heading(doc, "2. Histórico Clínico Detalhado")

    # Diagnósticos
    diags = f.getlist("diagnosticos")
    if diags:
        add_heading(doc, "Diagnósticos", level=2)
        for d in diags:
            label = DIAGNOSTICOS_LABELS.get(d, d)
            add_bullet(doc, label)

    # Detalhamento Fibromialgia
    if "fibromialgia" in diags:
        add_heading(doc, "Fibromialgia — Detalhamento", level=2)
        tempo   = f.get("fibro_tempo_diagnostico", "")
        intens  = f.get("fibro_intensidade_dor", "")
        crit    = f.get("fibro_criterios", "")
        sintomas_fibro = f.getlist("fibro_sintomas")

        if tempo:
            add_bullet(doc, tempo, "Tempo de diagnóstico:")
        if intens:
            add_bullet(doc, FIBRO_INTENSIDADE_LABELS.get(intens, intens), "Intensidade da dor (EVA):")
        if crit:
            add_bullet(doc, FIBRO_CRITERIOS_LABELS.get(crit, crit), "Critérios diagnósticos:")
        if sintomas_fibro:
            p = doc.add_paragraph()
            r = p.add_run("Sintomas relatados:")
            set_font(r, size=10.5, bold=True)
            for s in sintomas_fibro:
                add_bullet(doc, FIBRO_SINTOMAS_LABELS.get(s, s))

    # Antecedentes
    pre_natal = f.get("pre_natal", "")
    parto     = f.get("parto", "")
    hipoxia   = f.get("hipoxia", "")
    id_sint   = f.get("idade_sintomas", "")

    ante_items = []
    if pre_natal == "sim":   ante_items.append("Realizou pré-natal")
    elif pre_natal == "nao": ante_items.append("Não realizou pré-natal")
    if parto == "normal":          ante_items.append("Parto normal")
    elif parto == "cesareo":       ante_items.append("Parto cesáreo")
    elif parto == "cesareo_complicacoes": ante_items.append("Parto cesáreo com complicações")
    if hipoxia == "sim":   ante_items.append("Hipoxia neonatal: sim")
    elif hipoxia == "nao": ante_items.append("Hipoxia neonatal: não")
    if id_sint: ante_items.append(f"Idade dos primeiros sintomas: {id_sint}")

    if ante_items:
        add_heading(doc, "Antecedentes", level=2)
        for a in ante_items:
            add_bullet(doc, a)

    # Sintomas observados
    sintomas = f.getlist("sintomas")
    if sintomas:
        add_heading(doc, "Sintomas Observados", level=2)
        for s in sintomas:
            add_bullet(doc, SINTOMAS_LABELS.get(s, s))

    # Crises convulsivas
    tem_crises  = f.get("tem_crises", "")
    freq_crises = f.get("freq_crises", "")
    if tem_crises:
        add_heading(doc, "Crises Convulsivas", level=2)
        label_crise = "Sim" if tem_crises == "sim" else "Não"
        add_bullet(doc, label_crise, "Apresenta crises:")
        if freq_crises and tem_crises == "sim":
            add_bullet(doc, freq_crises, "Frequência:")

    # Escolaridade
    escola = f.get("escola", "")
    if escola:
        add_heading(doc, "Escolaridade", level=2)
        add_body(doc, escola, indent=True)

    # Notas do histórico
    hist_notas = f.get("historico_notas", "").strip()
    if hist_notas:
        add_heading(doc, "Observações do Histórico", level=2)
        add_body(doc, hist_notas, indent=True)

    # ════════ 3. TRATAMENTOS CONVENCIONAIS ════════
    add_heading(doc, "3. Tratamentos Convencionais Já Realizados")

    # Medicações
    meds_usados = []
    for key, nome_med in MEDS:
        if f.get(f"med_{key}_usou"):
            dose     = f.get(f"med_{key}_dose", "").strip()
            resposta = f.get(f"med_{key}_resposta", "")
            linha = nome_med
            if dose:
                linha += f" — {dose}"
            if resposta:
                linha += f" ({MED_RESPOSTA_LABELS.get(resposta, resposta)})"
            meds_usados.append(linha)

    outras_meds = f.get("med_outras", "").strip()

    if meds_usados or outras_meds:
        add_heading(doc, "Medicações Utilizadas", level=2)
        for m in meds_usados:
            add_bullet(doc, m)
        if outras_meds:
            for linha in outras_meds.splitlines():
                if linha.strip():
                    add_bullet(doc, linha.strip())

    # Terapias
    terapias = f.getlist("terapias")
    if terapias:
        add_heading(doc, "Terapias Multidisciplinares", level=2)
        for t in terapias:
            add_bullet(doc, TERAPIAS_LABELS.get(t, t))

    # CBD prévio
    cbd_previo = f.get("cbd_previo", "")
    cbd_dur    = f.get("cbd_duracao", "").strip()
    cbd_motivo = f.get("cbd_motivo_interrupcao", "").strip()
    if cbd_previo:
        add_heading(doc, "Uso Prévio de Canabidiol", level=2)
        cbd_label = {
            "sim": "Sim, com melhora",
            "sim_sem_resposta": "Sim, sem resposta terapêutica",
            "nao": "Não utilizou anteriormente",
        }.get(cbd_previo, cbd_previo)
        add_bullet(doc, cbd_label, "CBD prévio:")
        if cbd_dur:    add_bullet(doc, cbd_dur, "Duração:")
        if cbd_motivo: add_bullet(doc, cbd_motivo, "Motivo da interrupção:")

    # Notas tratamentos
    trat_notas = f.get("tratamentos_notas", "").strip()
    if trat_notas:
        add_body(doc, trat_notas, indent=True)

    # ════════ 4. JUSTIFICATIVA CLÍNICA ════════
    just_vals = f.getlist("justificativas")
    just_notas = f.get("justificativa_notas", "").strip()
    if just_vals or just_notas:
        add_heading(doc, "4. Justificativa Clínica para o Canabidiol")
        for j in just_vals:
            texto = JUSTIFICATIVAS_TEXTOS.get(j, "")
            if texto:
                add_body(doc, texto, indent=True)
                doc.add_paragraph()
        if just_notas:
            add_body(doc, just_notas, indent=True)

    # ════════ 5. ESPECIFICAÇÃO TÉCNICA ════════
    produto_v   = f.get("produto", "")
    produto_c   = f.get("produto_custom", "").strip()
    produto_nome = produto_c if produto_v == "custom" else PRODUTOS_LABELS.get(produto_v, "")

    if produto_nome:
        add_heading(doc, "5. Especificação Técnica do Produto")
        add_bullet(doc, produto_nome, "Produto prescrito:")

    # ════════ 6. POSOLOGIA ════════
    gotas       = f.get("pos_gotas", "")
    frequencia  = f.get("pos_frequencia", "")
    via         = f.get("pos_via", "")
    segundos    = f.get("pos_segundos", "")
    frascos     = f.get("pos_frascos", "")
    refeicao_v  = f.get("pos_refeicao", "")
    pos_obs     = f.get("pos_obs", "").strip()

    refeicao_labels = {
        "apos_cafe_jantar": "após café da manhã e jantar",
        "apos_refeicoes":   "após as refeições",
        "em_jejum":         "em jejum",
    }

    if gotas:
        add_heading(doc, "6. Posologia Prescrita")
        via_label = "sublingual" if via == "sublingual" else "oral"
        pos_text = f"{gotas} gota(s) {frequencia}, via {via_label}"
        if via == "sublingual" and segundos:
            pos_text += f", mantendo sublingual por {segundos} segundos"
        if refeicao_v:
            pos_text += f", {refeicao_labels.get(refeicao_v, '')}"
        pos_text += "."
        add_body(doc, pos_text, indent=True)

        if frascos:
            add_bullet(doc, f"{frascos} frasco(s)", "Quantidade:")
        if pos_obs:
            add_body(doc, pos_obs, indent=True)

    # ════════ 7. URGÊNCIA / RISCO ════════
    urgencias  = f.getlist("urgencias")
    urg_notas  = f.get("urgencia_notas", "").strip()
    if urgencias or urg_notas:
        add_heading(doc, "7. Urgência / Risco à Saúde")
        for u in urgencias:
            texto = URGENCIAS_TEXTOS.get(u, "")
            if texto:
                add_body(doc, texto, indent=True)
                doc.add_paragraph()
        if urg_notas:
            add_body(doc, urg_notas, indent=True)

    # ════════ 8. ACOMPANHAMENTO ════════
    acomp        = f.getlist("acompanhamento")
    periodicidade = f.get("acomp_periodicidade", "")
    acomp_notas  = f.get("acomp_notas", "").strip()

    if acomp or periodicidade or acomp_notas:
        add_heading(doc, "8. Declaração de Acompanhamento")
        if acomp:
            p = doc.add_paragraph()
            r = p.add_run("O acompanhamento do paciente será realizado com:")
            set_font(r, size=10.5)
            for a in acomp:
                add_bullet(doc, ACOMP_LABELS.get(a, a))
        if periodicidade:
            add_bullet(doc, PERIOD_LABELS.get(periodicidade, periodicidade), "Periodicidade:")
        if acomp_notas:
            add_body(doc, acomp_notas, indent=True)

    # ════════ CIDs ════════
    cids     = f.getlist("cids")
    cid_extra = f.get("cid_extra", "").strip()

    CID_LABELS = {
        "F84.0": "F84.0 — Autismo infantil",
        "F84.1": "F84.1 — Autismo atípico",
        "F84.5": "F84.5 — Síndrome de Asperger",
        "F90.0": "F90.0 — Transtorno de hiperatividade com déficit de atenção",
        "G40.0": "G40.0 — Epilepsia e síndromes epilépticas idiopáticas",
        "G40.2": "G40.2 — Epilepsia e síndromes epilépticas sintomáticas",
        "G40.3": "G40.3 — Epilepsia generalizada",
        "F70":   "F70 — Retardo mental leve",
        "F71":   "F71 — Retardo mental moderado",
        "F80":   "F80 — Transtornos específicos do desenvolvimento da fala",
        "F51.0": "F51.0 — Insônia não-orgânica",
        "M79.7": "M79.7 — Fibromialgia",
    }

    if cids or cid_extra:
        add_heading(doc, "CID — Classificação Internacional de Doenças")
        for c in cids:
            add_bullet(doc, CID_LABELS.get(c, c))
        if cid_extra:
            add_bullet(doc, cid_extra)

    # ════════ ASSINATURA ════════
    doc.add_paragraph()
    add_divider(doc)
    doc.add_paragraph()

    cidade = f.get("cidade", "")
    data_v = f.get("data", "")
    if data_v:
        try:
            dt = datetime.strptime(data_v, "%Y-%m-%d")
            data_fmt = dt.strftime("%d de %B de %Y").replace(
                "January","janeiro").replace("February","fevereiro").replace(
                "March","março").replace("April","abril").replace(
                "May","maio").replace("June","junho").replace(
                "July","julho").replace("August","agosto").replace(
                "September","setembro").replace("October","outubro").replace(
                "November","novembro").replace("December","dezembro")
        except Exception:
            data_fmt = data_v
    else:
        data_fmt = ""

    local_data = ""
    if cidade and data_fmt:
        local_data = f"{cidade}, {data_fmt}"
    elif cidade:
        local_data = cidade
    elif data_fmt:
        local_data = data_fmt

    if local_data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(local_data)
        set_font(r, size=10.5)

    doc.add_paragraph()
    doc.add_paragraph()

    p_ass = doc.add_paragraph()
    p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ass = p_ass.add_run("_" * 45)
    set_font(r_ass, size=10.5)

    p_nome = doc.add_paragraph()
    p_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_nome = p_nome.add_run(f"Dr(a). {medico}")
    set_font(r_nome, size=10.5, bold=True)

    p_crm = doc.add_paragraph()
    p_crm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crm_line = f"CRM {crm}"
    if especialidade:
        crm_line += f" | {especialidade}"
    r_crm = p_crm.add_run(crm_line)
    set_font(r_crm, size=10)

    return doc


# ── Rotas ─────────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/gerar", methods=["POST"])
def gerar():
    f = request.form

    logo_bytes = None
    logo_file = request.files.get("logo")
    if logo_file and logo_file.filename:
        logo_bytes = logo_file.read()

    doc = gerar_docx(f, logo_bytes)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    nome_paciente = f.get("paciente_nome", "paciente").replace(" ", "_")
    filename = f"relatorio_{nome_paciente}.docx"

    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    import os
    port = int(os.environ.get("PORT", 5005))
    app.run(host="0.0.0.0", port=port)
